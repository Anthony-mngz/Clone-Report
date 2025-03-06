import yaml
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Inches, Emu
import pandas as pd


def load_yaml_config(yaml_path):
    """Charge le fichier de configuration YAML."""
    with open(yaml_path, "r", encoding="utf-8") as file:
        return yaml.safe_load(file)


def apply_format(run, format_config):
    for attr, value in format_config.items():
        if attr == "font":
            run.font.name = value
        elif attr == "size":
            run.font.size = Pt(value)
        elif attr == "bold":
            run.bold = value


def replace_text(paragraph, replacements, format_config=None):
    for key, value in replacements.items():
        if isinstance(value, (int, float)) and "rounding" in replacements:
            value = round(value, replacements["rounding"])
        if key in paragraph.text:
            paragraph.text = paragraph.text.replace(key, str(value))
            if format_config:
                for run in paragraph.runs:
                    apply_format(run, format_config)


def insert_image(doc, image_config):
    if image_config:
        image_path = image_config.get("path")
        size = image_config.get("size")
        if image_path:
            if size:
                width, height = size
                doc.add_picture(image_path, width=Inches(width), height=Inches(height))
            else:
                doc.add_picture(image_path)


def insert_table(doc, table_config):
    if not table_config:
        print("Données de tableau vides, table ignorée.")
        return

    data = table_config.get("data", [])
    headers = table_config.get("headers", [])

    if not data:
        print("Données vides pour la table.")
        return

    if not list(doc.sections):
        print("Aucune section détectée, ajout d'une section par défaut.")
        doc.add_section()

    section = doc.sections[-1]
    if section.page_width is None or section.left_margin is None or section.right_margin is None:
        print("Certaines propriétés de la section sont absentes. Définition de valeurs par défaut.")
        section.page_width = Emu(12240)  # Valeur par défaut pour la largeur de page A4
        section.left_margin = Emu(1440)  # Valeur par défaut pour la marge gauche (2,54 cm)
        section.right_margin = Emu(1440)  # Valeur par défaut pour la marge droite (2,54 cm)

    try:
        table = doc.add_table(rows=1, cols=len(headers), style="Table Grid")
    except KeyError:
        print("Style 'Table Grid' introuvable. Utilisation du style par défaut.")
        table = doc.add_table(rows=1, cols=len(headers))

    hdr_cells = table.rows[0].cells
    for idx, header in enumerate(headers):
        hdr_cells[idx].text = str(header)

    for row in data:
        row_cells = table.add_row().cells
        for idx, cell in enumerate(row):
            row_cells[idx].text = str(cell)

    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    apply_format(run, table_config.get("format", {}))


def replace_text_in_docx(doc_path, yaml_path, output_path):
    config = load_yaml_config(yaml_path)
    doc = Document(doc_path)

    if not list(doc.sections):
        print("Aucune section détectée, ajout d'une section par défaut.")
        doc.add_section()

    if not doc.paragraphs:
        doc.add_paragraph("Section vide initialisée")

    for paragraph in doc.paragraphs:
        if "text" in config:
            replace_text(paragraph, config["text"], config.get("format"))

    if "image" in config:
        insert_image(doc, config["image"])

    if "table" in config:
        insert_table(doc, config["table"])

    if "array_from_excel" in config:
        df = pd.read_excel(config["array_from_excel"]["path"], engine="openpyxl")
        table_config = {"data": df.values.tolist(), "headers": df.columns.tolist(), "format": config.get("format", {})}
        insert_table(doc, table_config)

    doc.save(output_path)
    print(f"Document modifié enregistré sous : {output_path}")


# Exemple d'utilisation
doc_path = "template.docx"
yaml_path = "config.yaml"
output_path = "output.docx"
replace_text_in_docx(doc_path, yaml_path, output_path)
