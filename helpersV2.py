import yaml
from docx import Document
from docx.shared import Pt
from docx.shared import Inches
import pandas as pd
import os

def load_yaml_config(yaml_path):
    """Charge le fichier de configuration YAML avec gestion des erreurs."""
    try:
        with open(yaml_path, "r", encoding="utf-8") as file:
            return yaml.safe_load(file)
    except Exception as e:
        print(f"Erreur lors du chargement du fichier YAML : {e}")
        raise

def apply_format(run, format_config):
    """Applique le formatage sur le texte (police, taille, gras)."""
    if "font" in format_config:
        run.font.name = format_config["font"]
    if "size" in format_config:
        run.font.size = Pt(format_config["size"])
    if "bold" in format_config:
        run.bold = format_config["bold"]

def replace_text(paragraph, replacements):
    """Remplace le texte dans un paragraphe selon les configurations YAML."""
    for key, value in replacements.items():
        if isinstance(value, (int, float)) and "rounding" in replacements:
            value = round(value, replacements["rounding"])
        if key in paragraph.text:
            paragraph.text = paragraph.text.replace(key, str(value))

def insert_image(doc, image_path, size):
    """Insère une image dans le document avec la taille donnée."""
    if not os.path.exists(image_path):
        print(f"Erreur : l'image à {image_path} n'existe pas.")
        return
    if size:
        width, height = size
        doc.add_picture(image_path, width=Inches(width), height=Inches(height))
    else:
        doc.add_picture(image_path)

def insert_table(doc, data):
    """Insère un tableau dans le document à partir des données."""
    if not data:
        print("Données de tableau vides, table ignorée.")
        return

    # Ajoute une section vide si le document est vide
    if not list(doc.sections):
        doc.add_section()

    section = doc.sections[-1]
    if section.page_width is None:
        section.page_width = Inches(8.5)  # Largeur par défaut A4
    if section.left_margin is None:
        section.left_margin = Inches(1)  # Marge gauche par défaut
    if section.right_margin is None:
        section.right_margin = Inches(1)  # Marge droite par défaut

    if len(doc.paragraphs) == 0:
        doc.add_paragraph("Section vide initialisée")

    # Création de la table avec en-têtes
    table = doc.add_table(rows=1, cols=len(data[0]))

    hdr_cells = table.rows[0].cells

    # Remplissage des en-têtes
    for col_idx, header in enumerate(data[0].keys()):
        hdr_cells[col_idx].text = str(header)

    # Remplissage des lignes
    for row in data:
        row_cells = table.add_row().cells
        for col_idx, cell in enumerate(row.values()):
            row_cells[col_idx].text = str(cell)

    if len(data) > 1:  # Si plusieurs lignes, on ajoute une ligne vide
        table.add_row()
        row_cells = table.rows[-1].cells
        for col_idx, cell in enumerate(data[-1].values()):
            row_cells[col_idx].text = str(cell)

def insert_table_from_excel(doc, excel_path):
    """Insère un tableau depuis un fichier Excel."""
    try:
        df = pd.read_excel(excel_path)
        insert_table(doc, df.to_dict(orient="records"))
    except Exception as e:
        print(f"Erreur lors de l'importation du fichier Excel : {e}")

def replace_text_in_docx(doc_path, yaml_path, output_path):
    """Remplace le texte dans un document Word en fonction d'un fichier de configuration YAML."""
    config = load_yaml_config(yaml_path)
    doc = Document(doc_path)

    # Vérification si le document est vide, sinon ajout d'une section et d'un paragraphe
    if not list(doc.sections):
        doc.add_section()

    if not doc.paragraphs:
        doc.add_paragraph("Section vide initialisée")

    # Remplacement du texte dans les paragraphes
    for paragraph in doc.paragraphs:
        if "text" in config:
            replace_text(paragraph, config["text"])
            if "format" in config:
                for run in paragraph.runs:
                    apply_format(run, config["format"])

    # Insertion de l'image
    if "image" in config:
        insert_image(doc, config["image"]["path"], config["image"].get("size"))

    # Insertion du tableau depuis les données YAML
    if "table" in config:
        headers = config["table"].get("headers", [])
        data = config["table"].get("data", [])
        if headers and data:
            table_data = [dict(zip(headers, row)) for row in data]
            insert_table(doc, table_data)

    # Insertion d'un tableau depuis un fichier Excel
    if "array_from_excel" in config:
        excel_path = config["array_from_excel"]["path"]
        insert_table_from_excel(doc, excel_path)

    # Sauvegarde du document modifié
    try:
        doc.save(output_path)
        print(f"Document modifié enregistré sous : {output_path}")
    except Exception as e:
        print(f"Erreur lors de l'enregistrement du fichier : {e}")

# Exemple d'utilisation
doc_path = "template.docx"
yaml_path = "config.yaml"
output_path = "output.docx"
replace_text_in_docx(doc_path, yaml_path, output_path)
